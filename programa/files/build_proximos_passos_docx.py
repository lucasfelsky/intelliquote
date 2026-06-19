"""Gera o documento de proximos passos e implementacoes futuras do IntelliQuote."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\User\OneDrive - SQ Quimica\Área de Trabalho\PROJETOS\Intelliquote\documentação\Proximos_Passos_Implementacoes_IntelliQuote.docx"

ACCENT = RGBColor(0x18, 0x40, 0x54)        # navy
ACCENT_2 = RGBColor(0x00, 0xAE, 0x91)      # teal
MUTED = RGBColor(0x4A, 0x55, 0x60)
TEXT = RGBColor(0x1F, 0x29, 0x33)

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_run(p, text, *, bold=False, italic=False, color=None, size=11, font='Calibri'):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color if color else TEXT
    # forca a fonte para East Asian/complex script
    rpr = r._element.get_or_add_rPr()
    rfont = rpr.find(qn('w:rFonts'))
    if rfont is None:
        rfont = OxmlElement('w:rFonts')
        rpr.append(rfont)
    rfont.set(qn('w:ascii'), font)
    rfont.set(qn('w:hAnsi'), font)
    rfont.set(qn('w:cs'), font)
    rfont.set(qn('w:eastAsia'), font)
    return r

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, bold=True, color=ACCENT, size=20, font='Calibri')
    # linha decorativa
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '00AE91')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, color=ACCENT, size=14, font='Calibri')
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, color=ACCENT_2, size=12, font='Calibri')
    return p

def para(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=bold, italic=italic, color=color if color else TEXT, size=size)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    # limpa runs pre-existentes do estilo
    for r in list(p.runs):
        r.text = ''
    add_run(p, text)
    return p

def check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, '\u2610  ', color=ACCENT_2, bold=True)
    add_run(p, text)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, italic=True, color=MUTED, size=9)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def make_table(doc, headers, rows, col_widths=None, header_bg='184054', header_fg=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Light Grid Accent 1'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, color=header_fg, size=10)
    # body
    for r_idx, row in enumerate(rows):
        bg = 'F5F9F8' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            add_run(p, str(val), size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table


doc = Document()

# margens
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# style padrao
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============= CAPA =============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
add_run(p, 'IntelliQuote', bold=True, color=ACCENT, size=36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Plano de Próximos Passos e Implementações Futuras', bold=True, color=ACCENT_2, size=18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
add_run(p, 'Roadmap pós-feedback e pós-refactor visual', italic=True, color=MUTED, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
add_run(p, 'Versão 1.0   |   16/06/2026', color=MUTED, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
add_run(p, 'Documento de planejamento', color=MUTED, size=11)

page_break(doc)

# ============= SUMARIO =============
h1(doc, 'Sumário')
toc_items = [
    '1.  Contexto e ponto de partida',
    '2.  Princípios que guiam o roadmap',
    '3.  Fase 0 — Estabilização e hardening (próximas 2 semanas)',
    '4.  Fase 1 — Migração para Firebase Auth (Caminho C)',
    '5.  Fase 2 — Indicadores e dashboards',
    '6.  Fase 3 — Experiência do usuário e acessibilidade',
    '7.  Fase 4 — Integração com o Portal COMEX',
    '8.  Backlog técnico e de dívida',
    '9.  Marcos de qualidade e governança',
    '10. Riscos, mitigações e dependências externas',
    '11. Cronograma de alto nível',
    '12. Apêndice — referências e arquivos',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, item, size=11)

page_break(doc)

# ============= 1. CONTEXTO =============
h1(doc, '1. Contexto e ponto de partida')

para(doc,
     'Este documento descreve os próximos passos e as implementações planejadas para o IntelliQuote após o ciclo de correções disparado pelo feedback dos usuários (pacote de 05/05/2026) e o refactor de linguagem visual concluído em 16/06/2026. O objetivo é alinhar a evolução do produto, garantir que a fase atual permaneça estável e priorizar as frentes com maior retorno sobre o esforço.')

h2(doc, '1.1. Onde estamos hoje')
bullet(doc, 'Backend Node.js + Express + Prisma + PostgreSQL (Supabase) com 46 testes automatizados passando (1 ignorado) e 15 arquivos de teste cobrindo rotas, validadores, controllers, middlewares e fluxos de auditoria.')
bullet(doc, 'Frontend estático (HTML/CSS/JS) com 8 abas de workspace (Dashboard, Fornecedores, Cotações, Propostas, Comparações, Relatórios, Auditoria, Usuários/Ajuda).')
bullet(doc, 'Autenticação por JWT (access 15 min + refresh 7 d) em cookie HttpOnly, com recuperação de senha por token de uso único (TTL 1 h) e rate limit por IP.')
bullet(doc, 'Refactor visual concluído: paleta teal #00AE91 + navy #184054, fonte Roboto, cards com raio 22-32 px, botões em gradiente marinho, alinhado à linguagem do Portal COMEX.')
bullet(doc, 'Documentação interna: relatório de correções e melhorias (Relatorio_Correcoes_Melhorias_IntelliQuote.docx) e documento de decisão de camada de dados (DATA_LAYER_DECISION.md) já entregues.')

h2(doc, '1.2. Visão de produto')
para(doc,
     'Consolidar o IntelliQuote como a ferramenta padrão de cotações internacionais da SQ Química, integrada ao Portal COMEX, com autenticação unificada, indicadores de economia e lead time, e fluxos auditáveis fim a fim.')

# ============= 2. PRINCIPIOS =============
h1(doc, '2. Princípios que guiam o roadmap')

bullet(doc, 'Estabilidade primeiro: nenhuma nova funcionalidade que dependa de dados novos entra antes da fase de hardening.')
bullet(doc, 'Mudanças mensuráveis: cada fase tem indicador de sucesso quantitativo (latência, taxa de erro, NPS interno, cobertura de testes).')
bullet(doc, 'Compatibilidade com o COMEX: qualquer decisão visual, de auth ou de dados precisa conversar com a stack do Portal COMEX.')
bullet(doc, 'Reversibilidade: cada entrega mantém caminho de rollback documentado (ex.: styles.legacy.css, variáveis de ambiente por escopo).')
bullet(doc, 'Documentação junto do código: changelog atualizado a cada entrega; decisões relevantes viram documento em /programa/docs.')

# ============= 3. FASE 0 =============
h1(doc, '3. Fase 0 — Estabilização e hardening (próximas 2 semanas)')

para(doc,
     'A meta é travar a versão atual e reduzir a chance de regressões antes de iniciar a migração de auth. Tudo aqui é trabalho de manutenção, sem mudar contrato de API nem modelo de dados.')

h2(doc, '3.1. Cobertura de testes')
bullet(doc, 'Subir a cobertura para ≥ 80% em services/ e ≥ 70% em controllers/, com relatório via vitest --coverage no CI.')
bullet(doc, 'Adicionar testes de carga leves (autocannon) para forgot-password, login e POST /api/v1/quote-requests.')
bullet(doc, 'Criar teste de fumaça que sobe o app em test mode e percorre todas as 8 abas renderizadas (jsdom + supertest).')

h2(doc, '3.2. Observabilidade')
bullet(doc, 'Adicionar request-id em todas as respostas e correlacionar com logs estruturados (pino + pretty em dev, JSON em prod).')
bullet(doc, 'Expor /metrics em formato Prometheus (latência, taxa de erro por rota, tamanho de payload, contagem de anexos servidos).')
bullet(doc, 'Configurar alerta no Supabase para conexões próximas do limite do plano free.')

h2(doc, '3.3. Segurança')
bullet(doc, 'Rodar npm audit e corrigir vulnerabilidades altas/críticas; travar com dependabot.')
bullet(doc, 'Forçar HTTPS-only no cookie HttpOnly, SameSite=Lax e Secure em produção (atualmente é Lax).')
bullet(doc, 'Revisar a lista de middlewares: garantir helmet CSP, HSTS, X-Frame-Options DENY em produção.')
bullet(doc, 'Adicionar proteção contra CSRF para mutações em cookie de sessão (double-submit token).')

h2(doc, '3.4. Banco de dados')
bullet(doc, 'Adicionar índice composto em QuoteRequest (status, createdAt) e em QuoteResponse (quoteRequestId, totalPrice).')
bullet(doc, 'Documentar rotina de backup do Supabase (pg_dump diário + restore test mensal).')
bullet(doc, 'Revisar migrations pendentes e fechar cadeia de prisma migrate status antes da Fase 1.')

h2(doc, '3.5. Critério de saída da Fase 0')
make_table(doc,
    ['Indicador', 'Meta', 'Como medir'],
    [
        ['Cobertura services/', '≥ 80%', 'vitest --coverage'],
        ['Cobertura controllers/', '≥ 70%', 'vitest --coverage'],
        ['Vulnerabilidades npm audit', '0 altas/críticas', 'npm audit --omit=dev'],
        ['Latência p95 /health/ready', '< 80 ms', 'k6 smoke'],
        ['Build + testes no CI', 'verdes em 3 PRs consecutivas', 'GitHub Actions'],
    ],
    col_widths=[Cm(5.5), Cm(4.5), Cm(6.5)])

# ============= 4. FASE 1 =============
h1(doc, '4. Fase 1 — Migração para Firebase Auth (Caminho C)')

para(doc,
     'Conforme registrado em programa/docs/DATA_LAYER_DECISION.md, o Caminho C é o de menor risco e maior retorno no curto prazo: troca o adapter de autenticação para Firebase Auth no mesmo projeto do Portal COMEX, mantendo o Postgres atual. Esta fase dura entre 3 e 5 dias úteis e libera o login unificado.')

h2(doc, '4.1. Pré-requisitos')
bullet(doc, 'Projeto Firebase provisionado (pode ser o mesmo do Portal COMEX, em ambiente de dev separado).')
bullet(doc, 'Authentication habilitada com provider e-mail/senha.')
bullet(doc, 'Custom claims definidos: role (admin, gestor, comprador, viewer) e supplierScope (lista de supplierIds visíveis, opcional).')
bullet(doc, 'Variáveis de ambiente adicionadas ao .env.example: FIREBASE_PROJECT_ID, FIREBASE_CLIENT_EMAIL, FIREBASE_PRIVATE_KEY, FIREBASE_WEB_API_KEY.')

h2(doc, '4.2. Backend')
bullet(doc, 'Adicionar firebase-admin como dependência.')
bullet(doc, 'Criar src/middleware/requireFirebaseAuth.ts validando Authorization: Bearer <idToken> via firebase-admin/auth.verifyIdToken.')
bullet(doc, 'Manter o cookie HttpOnly para o frontend estático: novo endpoint POST /api/v1/auth/session troca o id token por cookie de sessão.')
bullet(doc, 'Substituir o middleware atual de forma gradual com feature flag (USE_FIREBASE_AUTH=true|false).')
bullet(doc, 'Atualizar /api/v1/auth/forgot-password para usar firebase/auth.generatePasswordResetLink (mantém fallback local via feature flag).')
bullet(doc, 'Atualizar src/services/audit.ts para registrar actorUid e actorEmail vindos do token.')

h2(doc, '4.3. Frontend')
bullet(doc, 'Carregar firebase JS SDK via CDN no index.html (mesma versão do Portal COMEX).')
bullet(doc, 'Criar public/firebase-init.js com config e export de signIn/signOut/onAuthStateChanged.')
bullet(doc, 'Substituir o form de login atual por signInWithEmailAndPassword e enviar o idToken no header de toda requisição.')
bullet(doc, 'Tratativa de expiração: SDK renova automaticamente; onAuthStateChanged dispara novo cookie em /api/v1/auth/session.')

h2(doc, '4.4. Migração de dados')
bullet(doc, 'Script de bootstrap que, no primeiro login, cria/sincroniza o registro local em users a partir dos custom claims.')
bullet(doc, 'Se o e-mail não existir no Firebase, criar via firebase-admin/auth.createUser com senha temporária enviada por e-mail.')
bullet(doc, 'Mapear role legado (campo role na tabela users) para custom claim; remover do payload do token depois de migrado.')

h2(doc, '4.5. Testes e rollout')
bullet(doc, 'Adicionar testes do middleware requireFirebaseAuth (token válido, expirado, sem role, com role inválido).')
bullet(doc, 'Adicionar teste E2E: login via SDK JS, chamada autenticada, logout.')
bullet(doc, 'Rollout em ambiente de homologação por 1 semana; corte em produção em horário de baixo uso.')
bullet(doc, 'Manter USE_FIREBASE_AUTH=false no .env de produção por 1 semana após o corte, para rollback rápido.')

h2(doc, '4.6. Critério de saída da Fase 1')
make_table(doc,
    ['Indicador', 'Meta', 'Como medir'],
    [
        ['Logins unificados com COMEX', '100% em prod', 'audit log + dashboard'],
        ['Tempo médio de login', '< 1,5 s', 'log estruturado'],
        ['Tokens JWT próprios emitidos', '0', 'grep nos logs'],
        ['Testes automatizados', 'verdes', 'vitest run'],
        ['Rollback executado', '0 vezes na primeira semana', 'incident log'],
    ],
    col_widths=[Cm(5.5), Cm(4.5), Cm(6.5)])

# ============= 5. FASE 2 =============
h1(doc, '5. Fase 2 — Indicadores e dashboards')

h2(doc, '5.1. KPIs prioritários')
bullet(doc, 'Economia estimada por cotação adjudicada (R$ e %) — já existe no módulo de relatórios, precisa de tendência temporal.')
bullet(doc, 'Lead time médio por fornecedor e por Incoterm (FOB, CIF, EXW etc.).')
bullet(doc, 'Taxa de adjudicação (propostas com isWinner=true / total de respostas).')
bullet(doc, 'Tempo médio de resposta dos fornecedores (criação da cotação → primeira resposta).')
bullet(doc, 'Taxa de conclusão do onboarding em 3 passos (meta ≥ 90% em 30 dias).')

h2(doc, '5.2. Implementação')
bullet(doc, 'Criar endpoint GET /api/v1/reports/kpi com séries temporais (granularidade dia/semana/mês).')
bullet(doc, 'Adicionar widget na aba Dashboard consumindo o endpoint, com sparkline e comparação ao período anterior.')
bullet(doc, 'Exportar CSV/XLSX para os mesmos indicadores (com rate limit).')
bullet(doc, 'Criar view materializada report_kpi_daily refrescada por cron (a cada 15 min).')

h2(doc, '5.3. Critério de saída da Fase 2')
make_table(doc,
    ['Indicador', 'Meta'],
    [
        ['Widgets de KPI no Dashboard', '5 widgets'],
        ['Exportação CSV/XLSX', 'disponível para todos os relatórios'],
        ['Latência p95 /reports/kpi', '< 250 ms'],
        ['Adesão semanal ao dashboard', '≥ 60% dos usuários ativos'],
    ],
    col_widths=[Cm(9.0), Cm(7.5)])

# ============= 6. FASE 3 =============
h1(doc, '6. Fase 3 — Experiência do usuário e acessibilidade')

h2(doc, '6.1. UX e fluxo')
bullet(doc, 'Mover navegação por abas do topo para uma sidebar lateral recolhível (consistência com o Portal COMEX). Manter as tabs como secundárias no mobile.')
bullet(doc, 'Adicionar atalhos de teclado: / foca a busca global, g+d vai para Dashboard, g+s vai para Suppliers, g+q vai para Quote Requests.')
bullet(doc, 'Adicionar busca global (Cmd/Ctrl+K) cruzando fornecedores, cotações e propostas.')
bullet(doc, 'Implementar "salvar como rascunho" em todas as abas com formulários longos (Propostas, Fornecedores, Cotações).')

h2(doc, '6.2. Acessibilidade')
bullet(doc, 'Auditoria WCAG 2.1 AA com axe-core e correção dos achados (contraste, foco visível, aria-labels).')
bullet(doc, 'Adicionar prefers-reduced-motion para desabilitar animações de hover/transição.')
bullet(doc, 'Garantir navegação completa por teclado em todos os fluxos (especialmente wizard de 4 etapas).')
bullet(doc, 'Adicionar suporte a leitor de tela nas tabelas grandes (propostas, audit log) com aria-sort e aria-rowcount.')

h2(doc, '6.3. Mobile e PWA')
bullet(doc, 'Transformar o frontend em PWA instalável (manifest.json + service worker).')
bullet(doc, 'Garantir layout responsivo nas 8 abas para larguras ≥ 360 px.')
bullet(doc, 'Offline-first para visualização (somente leitura) de cotações em andamento.')

h2(doc, '6.4. Critério de saída da Fase 3')
make_table(doc,
    ['Indicador', 'Meta'],
    [
        ['Achados axe-core', '0 violações críticas'],
        ['Contraste mínimo', '4,5:1 em todos os textos'],
        ['Navegação por teclado', '100% dos fluxos críticos'],
        ['Lighthouse PWA score', '≥ 90'],
        ['NPS interno (time interno)', '≥ 8/10'],
    ],
    col_widths=[Cm(9.0), Cm(7.5)])

# ============= 7. FASE 4 =============
h1(doc, '7. Fase 4 — Integração com o Portal COMEX')

h2(doc, '7.1. Estratégia de integração')
para(doc,
     'A integração com o Portal COMEX é o motivador principal da migração para o Firebase. O escopo exato ainda precisa ser confirmado com o time do COMEX; o documento DATA_LAYER_DECISION.md trata os caminhos técnicos. Aqui listamos os pontos de produto que essa fase precisa endereçar.')

h2(doc, '7.2. Funcionalidades candidatas')
bullet(doc, 'Login único via Firebase Auth (Fase 1).')
bullet(doc, 'Notificações cruzadas: quando uma cotação é criada no IntelliQuote, o COMEX recebe um aviso; quando um processo do COMEX envolve importação, o IntelliQuote recebe um card de sugestão.')
bullet(doc, 'Cadastro de fornecedores compartilhado: o fornecedor criado em um sistema aparece no outro (com shadow write via Cloud Function).')
bullet(doc, 'Dashboard consolidado: KPIs do IntelliQuote embarcados no Portal COMEX via iframe autenticado.')
bullet(doc, 'Pesquisa de produtos: termo lançado no COMEX consulta a base de cotações históricas do IntelliQuote e sugere faixa de preço referência.')

h2(doc, '7.3. Pré-requisitos')
bullet(doc, 'Fase 0 e Fase 1 concluídas.')
bullet(doc, 'Reunião de kickoff com o time do COMEX para confirmar escopo e contratos de dados.')
bullet(doc, 'Cloud Functions de sincronização (IntelliQuote → COMEX e vice-versa) com idempotência e retry.')
bullet(doc, 'Política de privacidade atualizada (LGPD) cobrindo o compartilhamento de dados entre os dois produtos.')

h2(doc, '7.4. Critério de saída da Fase 4')
make_table(doc,
    ['Indicador', 'Meta'],
    [
        ['Funcionalidades integradas em prod', '≥ 2 (login único + 1 fluxo de dados)'],
        ['Latência de sincronização', '< 5 s ponta a ponta'],
        ['Incidentes P1 em 30 dias pós-lançamento', '0'],
        ['Conformidade LGPD', 'parecer jurídico favorável'],
    ],
    col_widths=[Cm(9.0), Cm(7.5)])

# ============= 8. BACKLOG =============
h1(doc, '8. Backlog técnico e de dívida')

para(doc, 'Itens já identificados que precisam ser endereçados em algum momento, sem data firme ainda. Devem ser revisitados a cada entrega.')

h3(doc, 'Backend')
bullet(doc, 'Mover cálculos de custo landed e comparação para services/ dedicados (já parcialmente feito; consolidar nomenclatura).')
bullet(doc, 'Substituir validações manuais repetidas em controllers por Zod schemas com .parse() em todos os endpoints — há rotas com validação parcial.')
bullet(doc, 'Introduzir fila assíncrona (Bull/BullMQ + Redis) para geração de relatórios pesados e envio de e-mails transacionais.')
bullet(doc, 'Padronizar respostas de erro com Problem Details (RFC 7807).')

h3(doc, 'Frontend')
bullet(doc, 'Quebrar app.js (atualmente > 1.900 linhas) em módulos ES com import/export e tree-shaking.')
bullet(doc, 'Trocar manipulação direta de DOM por um framework leve (Preact ou Alpine.js) para reduzir complexidade da aba Propostas.')
bullet(doc, 'Centralizar fetch em uma camada api.js com tratamento de 401/403/500 padronizado e refresh automático.')

h3(doc, 'Banco de dados')
bullet(doc, 'Avaliar separação de anexos (Attachment) em schema próprio + storage S3-compatible (ou Firebase Storage, alinhado à Fase 4).')
bullet(doc, 'Adicionar soft delete em QuoteRequest e Supplier (deletedAt) em vez de exclusão física.')
bullet(doc, 'Documentar convenções de nomenclatura (snake_case em SQL, camelCase em JS, PascalCase em classes).')

h3(doc, 'DevOps')
bullet(doc, 'Migrar CI do GitHub Actions para incluir deploy de preview por PR (Render ou Fly.io).')
bullet(doc, 'Criar ambiente de staging idêntico ao de produção (mesma versão do Postgres, mesmo bucket de anexos).')
bullet(doc, 'Adicionar scan de segredos (gitleaks) em pre-commit e no CI.')

# ============= 9. MARCOS =============
h1(doc, '9. Marcos de qualidade e governança')

h2(doc, '9.1. Definition of Done por entrega')
bullet(doc, 'Código revisado por outro dev (mesmo que informal).')
bullet(doc, 'Testes automatizados atualizados, build verde.')
bullet(doc, 'CHANGELOG.md atualizado com a data e o motivo.')
bullet(doc, 'Variáveis de ambiente documentadas em .env.example.')
bullet(doc, 'Smoke test manual em homologação (checklist no script smoke:local).')

h2(doc, '9.2. Cerimônias sugeridas')
bullet(doc, 'Weekly de produto (30 min): revisão do roadmap, próximos marcos, bloqueios.')
bullet(doc, 'Bi-weekly técnica (45 min): revisão de incidentes, débitos, decisões de arquitetura.')
bullet(doc, 'Mensal de governança (60 min): auditoria de acessos, revisão de LGPD, custo de infraestrutura.')

h2(doc, '9.3. Métricas globais de saúde do produto')
make_table(doc,
    ['Métrica', 'Meta', 'Frequência'],
    [
        ['Disponibilidade (uptime)', '≥ 99,5%', 'mensal'],
        ['Latência p95 nas rotas principais', '< 400 ms', 'semanal'],
        ['Cobertura de testes', '≥ 80% em services/', 'por entrega'],
        ['Vulnerabilidades altas/críticas', '0', 'por entrega'],
        ['NPS interno', '≥ 8/10', 'trimestral'],
        ['Custo de infraestrutura', '< US$ 50/mês', 'mensal'],
    ],
    col_widths=[Cm(6.5), Cm(4.5), Cm(5.5)])

# ============= 10. RISCOS =============
h1(doc, '10. Riscos, mitigações e dependências externas')

make_table(doc,
    ['Risco', 'Probabilidade', 'Impacto', 'Mitigação'],
    [
        ['Quebra de login durante migração para Firebase Auth', 'Média', 'Alto', 'Feature flag USE_FIREBASE_AUTH + cookie fallback; corte em horário de baixo uso'],
        ['Indisponibilidade do Supabase', 'Baixa', 'Alto', 'Backup diário + rotina de restore test; plano de contingência para Postgres self-hosted'],
        ['Sobrecarga de attachments no disco local', 'Média', 'Médio', 'Mover para S3/Firebase Storage; quotas por usuário'],
        ['Resistência do time à mudança de visual', 'Baixa', 'Médio', 'Manter styles.legacy.css; rollout gradual por aba; comunicação clara'],
        ['Escopo da integração com COMEX não fechado', 'Alta', 'Médio', 'Reunião de kickoff antes da Fase 4; congelar escopo da Fase 1 e 2'],
        ['Custo do Firestore crescer com o volume', 'Média', 'Médio', 'Rever antes da Fase 4; manter Supabase enquanto o ganho de integração não compensar'],
        ['Dependência de janela de manutenção do Portal COMEX', 'Média', 'Médio', 'Alinhar calendário com o time do COMEX na Fase 4'],
        ['Mudança de LGPD impactando o compartilhamento de dados', 'Baixa', 'Alto', 'Revisão jurídica antes da Fase 4; cláusula de minimização de dados'],
    ],
    col_widths=[Cm(4.5), Cm(2.5), Cm(2.5), Cm(7.0)])

# ============= 11. CRONOGRAMA =============
h1(doc, '11. Cronograma de alto nível')

para(doc, 'Estimativas em semanas úteis, a partir de 16/06/2026. Datas reais dependem de priorização e disponibilidade do time.')

make_table(doc,
    ['Fase', 'Janela', 'Duração estimada', 'Entregas principais'],
    [
        ['Fase 0', '16/06/2026 → 30/06/2026', '2 semanas', 'Cobertura de testes ≥ 80%, observabilidade, segurança, índices'],
        ['Fase 1', '01/07/2026 → 08/07/2026', '1 semana', 'Login unificado via Firebase Auth, feature flag, rollback'],
        ['Fase 2', '09/07/2026 → 29/07/2026', '3 semanas', 'KPIs no dashboard, exportação CSV/XLSX'],
        ['Fase 3', '30/07/2026 → 26/08/2026', '4 semanas', 'Sidebar, atalhos, busca global, acessibilidade, PWA'],
        ['Fase 4', 'a confirmar', 'a confirmar', 'Integração COMEX: notificações cruzadas, fornecedores compartilhados'],
    ],
    col_widths=[Cm(2.5), Cm(4.0), Cm(2.5), Cm(7.5)])

para(doc,
     'Observação: o cronograma presume 1 dev em tempo integral. Com 2 devs, as Fases 0 e 1 podem ser paralelizadas e a Fase 2 começa ainda em julho.',
     italic=True, color=MUTED, size=10)

# ============= 12. APENDICE =============
h1(doc, '12. Apêndice — referências e arquivos')

h2(doc, '12.1. Documentos relacionados')
bullet(doc, 'documentação/Feedback e Avaliação IntelliQuote.docx — fonte do ciclo de correções 05/05/2026.')
bullet(doc, 'documentação/Relatorio_Correcoes_Melhorias_IntelliQuote.docx — relatório final das 12 entregas do ciclo.')
bullet(doc, 'documentação/Proximos_Passos_Implementacoes_IntelliQuote.docx — este documento.')
bullet(doc, 'programa/docs/DATA_LAYER_DECISION.md — análise dos caminhos para Firebase.')
bullet(doc, 'programa/CHANGELOG.md — histórico de versões.')
bullet(doc, 'programa/README.md — guia de instalação e uso.')

h2(doc, '12.2. Pontos do código que serão tocados por fase')
make_table(doc,
    ['Fase', 'Arquivos previstos', 'Tipo de mudança'],
    [
        ['Fase 0', 'src/middleware/*, src/app.ts, prisma/schema.prisma', 'refactor + índices + cobertura'],
        ['Fase 1', 'src/middleware/, src/routes/auth.ts, src/services/audit.ts, public/firebase-init.js, public/app.js, public/index.html', 'troca de auth por Firebase'],
        ['Fase 2', 'src/controllers/report.ts, src/services/report.ts, public/app.js, public/index.html', 'novos endpoints + widgets'],
        ['Fase 3', 'public/styles.css, public/app.js, public/index.html, public/sw.js, public/manifest.json', 'UX, acessibilidade, PWA'],
        ['Fase 4', 'src/integrations/comex/*, firestore.rules, functions/*', 'integração entre produtos'],
    ],
    col_widths=[Cm(2.5), Cm(7.5), Cm(6.5)])

h2(doc, '12.3. Pessoas e contatos sugeridos')
bullet(doc, 'Product Owner do IntelliQuote: definir KPIs e prioridade de cada fase.')
bullet(doc, 'Tech Lead: revisar PRs de Fase 1 e Fase 4 (mudanças sensíveis).')
bullet(doc, 'Time do Portal COMEX: alinhar escopo da Fase 4 e janelas de manutenção.')
bullet(doc, 'Jurídico/Compliance: revisar cláusula de compartilhamento de dados antes da Fase 4.')

# rodape final
page_break(doc)
h1(doc, 'Encerramento')
para(doc,
     'Este roadmap é um documento vivo: deve ser revisado a cada entrega e ajustado conformefeedback dos usuários, mudanças no Portal COMEX e novas prioridades de negócio. A próxima revisão está prevista para a entrega da Fase 0 (30/06/2026).')

para(doc, '— Fim do documento —', italic=True, color=MUTED, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT_PATH)
print(f'Documento salvo em: {OUT_PATH}')
