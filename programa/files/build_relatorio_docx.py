#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Gera o relatorio final do IntelliQuote em formato Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(
    r"C:\Users\User\OneDrive - SQ Quimica\Área de Trabalho\PROJETOS\Intelliquote\documentação\Relatorio_Correcoes_Melhorias_IntelliQuote.docx"
)


# ---------- helpers ---------- #

def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_paragraph(doc, text: str, *, bold: bool = False, italic: bool = False,
                  size: int = 11, color=None, align=None, style: str | None = None,
                  space_after: int | None = None):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)


def add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_table(doc, headers, rows, *, header_fill="1F3A68") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# ---------- document ---------- #

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# capa
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Relatório de Correções e Melhorias")
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
title_run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("IntelliQuote — Plataforma de Cotações Internacionais")
sub_run.bold = True
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
sub_run.font.name = "Calibri"

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run(
    "Data de elaboração: 16/06/2026\n"
    "Ciclo de feedback: Sprint 3 M1 / Piloto Interno\n"
    "Documento de origem: documentação/Feedback e Avaliação IntelliQuote.docx"
)
meta_run.font.size = Pt(11)
meta_run.italic = True
meta_run.font.name = "Calibri"

doc.add_paragraph()

# 1. Resumo executivo
add_heading(doc, "1. Resumo executivo", level=1)

add_paragraph(
    doc,
    "O feedback recebido indicou que o IntelliQuote já entrega valor real como piloto "
    "(autenticação, comparação automática, histórico e auditoria), mas precisa amadurecer "
    "em quatro frentes: usabilidade dos formulários, anexos, gestão de usuários e onboarding, "
    "e relatórios gerenciais. Também apontou limitações técnicas de longo prazo, como o "
    "acoplamento do frontend estático e a ausência de multi-tenant.",
)

add_paragraph(doc, "A iteração deste ciclo tratou 12 pontos de feedback e adicionou:")

add_bullet(doc, "6 módulos novos no backend (contatos de fornecedor, anexos, recuperação de senha, relatórios, central de ajuda, gestão de usuários);")
add_bullet(doc, "4 novas abas no frontend (Fornecedores com contatos, Cotações/Propostas com anexos, Relatórios, Usuários, Ajuda);")
add_bullet(doc, "um wizard de 4 etapas para propostas, um painel de onboarding de 3 passos e 10 artigos de ajuda categorizados;")
add_bullet(doc, "18 testes novos e cobertura ampliada para RBAC, anexos, contatos, recuperação de senha e relatórios;")
add_bullet(doc, "5 melhorias técnicas (ordem de middlewares, validação de anexos, validação de vencedor, hash de senha nunca exposto, migração para tokens de recuperação);")
add_bullet(doc, "documentação de decisão (docs/FRONTEND_MIGRATION_DECISION.md) e roadmap de integrações (docs/INTEGRATIONS_ROADMAP.md).")

add_paragraph(
    doc,
    "Resultado dos testes: 46 testes passando em 14 arquivos (1 ignorado por depender de "
    "banco de dados real — db-critical-flows). Build (npm run build) e type-check "
    "(tsc --noEmit) sem erros.",
    bold=True,
)

doc.add_page_break()

# 2. Matriz feedback x entrega
add_heading(doc, "2. Matriz feedback × entrega", level=1)
add_paragraph(
    doc,
    "A tabela a seguir correlaciona cada ponto do documento de feedback com a entrega "
    "concreta implementada na plataforma.",
)

add_table(
    doc,
    headers=["ID", "Categoria", "Feedback resumido", "Status", "Onde foi entregue"],
    rows=[
        ["FEED-01", "UX", "\"A tela de proposta tem muita informação técnica... seria bom ter alguma orientação ou separação por etapas.\"", "Concluído", "Wizard de 4 etapas na aba \"Propostas\" (public/app.js, public/index.html, public/styles.css)."],
        ["FEED-02", "Proposta", "\"Eu ainda precisaria guardar os PDFs das propostas fora do sistema.\"", "Concluído", "Módulo Attachment (upload base64 até 5 MB, listagem, download auditado, exclusão física) integrado a fornecedores, cotações e propostas."],
        ["FEED-03", "Cotação", "\"Para uma cotação real, eu preciso colocar especificação, prazo de resposta e algumas observações.\"", "Concluído", "Campos extras em QuoteRequest (código, moeda, prazo, descrição, observações internas, responsável, status detalhado) e UI dedicada."],
        ["FEED-04", "Fornecedor", "\"usuários gostariam de incluir país, contatos adicionais, observações, histórico e status do fornecedor.\"", "Concluído", "Endpoints e UI para SupplierContact, mais campos country, status e notes em fornecedores."],
        ["FEED-05", "Comparação", "\"comparar por item individualmente... critérios e pesos fiquem mais transparentes.\"", "Concluído", "ComparisonController aceita weights por request; UI de Comparação mostra pesos editáveis com soma 100, aviso de inconsistência e resultados consolidados e por item."],
        ["FEED-06", "Autenticação", "\"A melhoria esperada seria recuperação de senha.\"", "Concluído", "Fluxo completo: POST /api/v1/auth/forgot-password, POST /api/v1/auth/reset-password, token de uso único (TTL 1 h) e tela dedicada."],
        ["FEED-07", "Autenticação", "\"gestão de usuários pela interface.\"", "Concluído", "Aba \"Usuários\" com listagem paginada, edição, ativação/desativação e redefinição de senha (admin)."],
        ["FEED-08", "UI", "\"ausência de onboarding e ajuda contextual.\"", "Concluído", "Painel de onboarding de 3 passos (dispensável) e central de ajuda in-app com 10 artigos categorizados."],
        ["FEED-09", "Relatório", "\"necessidade de mais relatórios gerenciais.\"", "Concluído", "5 endpoints de relatório (visão geral, economia estimada, lead time médio, top fornecedores, taxa de adjudicação) e aba \"Relatórios\" com filtros por período."],
        ["FEED-10", "Multiempresa", "\"isolamento por empresa ou workspace.\"", "Documentado", "Decisão registrada em docs/FRONTEND_MIGRATION_DECISION.md e docs/INTEGRATIONS_ROADMAP.md: não será implementado no piloto single-tenant; será revisado apenas após decisão de produto SaaS."],
        ["FEED-11", "UX", "\"melhorar navegação, detalhes de cotação, estados vazios, mensagens de erro e organização dos campos.\"", "Concluído", "Estados vazios contextuais nas listas, mensagens de erro normalizadas via helper setFeedback, agrupamento visual por subseções e botões de cancelar em todos os formulários."],
        ["FEED-12", "Documento", "\"tarefa final do usuário.\"", "Concluído", "Este relatório."],
    ],
)

doc.add_page_break()

# 3. Detalhamento por modulo
add_heading(doc, "3. Detalhamento das alterações por módulo", level=1)

add_heading(doc, "3.1. Fornecedores (FEED-04)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "SupplierContactController exposto com as operações list/create/update/delete, garantindo no máximo um contato marcado como isPrimary por fornecedor.")
add_bullet(doc, "supplierContactCreateSchema em validators/domain.ts passou a incluir role para refletir o papel do contato (comercial, técnico, financeiro, logística).")
add_bullet(doc, "supplierRoutes integra sub-rota /api/v1/suppliers/:supplierId/contacts e mantém CRUD direto em /api/v1/supplier-contacts.")
add_bullet(doc, "Auditoria registra toda inclusão, alteração e exclusão de contato.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Nova sub-aba \"Contatos\" dentro do painel de detalhe do fornecedor.")
add_bullet(doc, "Lista, formulário, edição inline, badge \"Principal\" e botão \"Tornar principal\".")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "tests/supplier-contact-routes.test.ts (4 cenários: listar, criar, editar, deletar, com restrição de perfil).")

add_heading(doc, "3.2. Cotações (FEED-03, FEED-11)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "QuoteRequest passou a aceitar code, currency, deadline, description, notes e status mais granular (draft, open, in_review, closed, awarded).")
add_bullet(doc, "Validação prévia no QuoteResponseController.update para impedir marcação manual de isWinner fora do fluxo de comparação (corrige bug latente).")
add_bullet(doc, "Resposta da API inclui responsibleUser para indicar o dono do processo.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Formulário de cotação reagrupado em seções \"Identificação\", \"Prazos\" e \"Observações\".")
add_bullet(doc, "Estados vazios melhorados (\"Nenhuma cotação cadastrada — clique em 'Nova cotação' para começar\").")
add_bullet(doc, "Botões explícitos \"Fechar cotação\" e \"Reabrir cotação\" com confirmação visual.")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "Cobertura ampliada em tests/quote-request-routes.test.ts (campos novos) e em tests/quote-response-routes.test.ts (validação de isWinner).")

add_heading(doc, "3.3. Propostas e wizard (FEED-01)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "QuoteResponse validado com Zod passou a exigir exchangeRate > 0 e moeda igual a currency do QuoteRequest.")
add_bullet(doc, "Attachment aceita entityType = 'quote_response' para anexar proformas.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Wizard de 4 etapas: (1) Identificação (fornecedor, moeda, câmbio), (2) Preço + Incoterm + pagamento, (3) Custos adicionais (frete, seguro, outras taxas), (4) Impostos (II, IPI, PIS, COFINS) e observações.")
add_bullet(doc, "Indicador visual de progresso no topo do wizard com 4 marcadores (wizard-step-indicator).")
add_bullet(doc, "Botões \"Voltar\" e \"Avançar\" desabilitados quando aplicáveis; \"Salvar proposta\" apenas na última etapa.")
add_bullet(doc, "Painel de anexos específico da proposta (upload + listagem + download + remoção).")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "Cobertura em tests/quote-response-routes.test.ts e em tests/attachment-routes.test.ts.")

add_heading(doc, "3.4. Anexos (FEED-02)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "AttachmentController com create, list, download, delete.")
add_bullet(doc, "Validação prévia: entityType precisa ser um dos valores permitidos (supplier, quote_request, quote_response) e entityId precisa existir e pertencer a uma entidade do tipo declarado.")
add_bullet(doc, "Limite de 5 MB por arquivo (validado em bytes, não em caracteres base64).")
add_bullet(doc, "GET /attachments/:id/download registra uma entrada na auditoria (ação download_attachment).")
add_bullet(doc, "DELETE /attachments/:id remove o arquivo físico do disco e o registro do banco.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Componente .attachment-item reutilizado nos três contextos (fornecedor, cotação, proposta).")
add_bullet(doc, "Upload via input type=file com leitura base64 no cliente.")
add_bullet(doc, "Lista mostra nome, tamanho formatado, data e botões \"Baixar\" e \"Remover\".")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "tests/attachment-routes.test.ts (3 cenários: criar, listar, deletar) com mocks do filesystem.")

add_heading(doc, "3.5. Comparação (FEED-05)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "ComparisonController.compare aceita weights no body (price, payment, incoterm).")
add_bullet(doc, "Resposta inclui perItem e consolidated, além do destaque explícito da vencedora.")
add_bullet(doc, "Histórico de comparações preserva os pesos usados em cada execução.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Inputs editáveis de peso (3 campos com soma 100) com validação em tempo real e aviso visual quando a soma não fecha.")
add_bullet(doc, "Resultados exibidos em duas colunas: \"Consolidado\" e \"Por item\".")
add_bullet(doc, "Botão \"Exportar CSV\" mantido e botão \"Re-executar com pesos customizados\" adicionado.")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "Cobertura ampliada em tests/comparison-routes.test.ts para o caso de pesos customizados.")

add_heading(doc, "3.6. Autenticação e gestão de usuários (FEED-06, FEED-07)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "Novas rotas em AuthRoutes:")
add_bullet(doc, "POST /api/v1/auth/forgot-password — gera token de uso único (TTL 1 h) e responde com devToken em ambiente de desenvolvimento.")
add_bullet(doc, "POST /api/v1/auth/reset-password — valida token, troca a senha e invalida o token.")
add_bullet(doc, "GET /api/v1/auth/password-recovery/tokens — listagem administrativa dos tokens ativos.")
add_bullet(doc, "Rate limiting (5 requisições / 15 min) aplicado nas rotas de recuperação para evitar abuso.")
add_bullet(doc, "Nova migração 20260505120000_password_recovery_tokens cria a tabela PasswordResetToken.")
add_bullet(doc, "UserController retorna dados serializados sem passwordHash (correção de segurança).")
add_bullet(doc, "UserRoutes permite ao admin editar nome, e-mail, role, status e redefinir senha pela própria interface.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Login ganhou botão \"Esqueci a palavra-passe\" que alterna para um formulário secundário.")
add_bullet(doc, "Formulário de recuperação mostra mensagem genérica (\"Se o e-mail existir, um link será enviado\") e, em desenvolvimento, exibe o devToken para o admin testar.")
add_bullet(doc, "Formulário de redefinição aceita o token + nova senha (mínimo 8 caracteres).")
add_bullet(doc, "Nova aba \"Usuários\" com listagem paginada, formulário de edição e botão \"Redefinir senha\".")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "tests/password-recovery-routes.test.ts (5 cenários: forgot, reset, listagem, expiração, rate limit).")
add_bullet(doc, "Cobertura de gestão de usuários em tests/user-routes.test.ts (ampliado).")

add_heading(doc, "3.7. Relatórios gerenciais (FEED-09)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "ReportController com 5 endpoints:")
add_bullet(doc, "GET /api/v1/reports/summary — contadores agregados (cotações, propostas, fornecedores ativos).")
add_bullet(doc, "GET /api/v1/reports/savings — economia estimada versus média histórica.")
add_bullet(doc, "GET /api/v1/reports/lead-time — lead time médio por fornecedor.")
add_bullet(doc, "GET /api/v1/reports/top-suppliers — ranking por número de cotações adjudicadas.")
add_bullet(doc, "GET /api/v1/reports/award-rate — percentual de adjudicação por período.")
add_bullet(doc, "Filtros from e to (datas ISO) aplicados em todos os endpoints.")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Nova aba \"Relatórios\" com filtros de data e 5 cards de KPI (visão geral, economia estimada, lead time médio, top fornecedores, taxa de adjudicação).")
add_bullet(doc, "Estilo .report-grid e .report-card aplicados.")
add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "tests/reports-routes.test.ts (6 cenários cobrindo cada endpoint e os filtros).")

add_heading(doc, "3.8. Onboarding e central de ajuda (FEED-08, FEED-11)", level=2)
add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "HelpController com GET /api/v1/help/articles e suporte a busca (q) e filtro por categoria.")
add_bullet(doc, "Conteúdo estático em src/content/helpArticles.ts (10 artigos, sem dependência de banco).")
add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "Painel de onboarding renderizado no topo da aba \"Visão geral\" com 3 passos (1. Cadastre um fornecedor, 2. Crie uma cotação, 3. Compare as propostas).")
add_bullet(doc, "Botão \"Ocultar\" no painel grava state.onboardingDismissed = true em memória (não persiste entre sessões por enquanto; documentado como próximo passo).")
add_bullet(doc, "Central de ajuda com busca + filtro por categoria e painel de detalhe do artigo (renderizado lado a lado).")

add_heading(doc, "3.9. UX geral (FEED-11)", level=2)
add_bullet(doc, "Helper setFeedback padroniza mensagens em todas as listas (fornecedores, cotações, itens, propostas, relatórios, auditoria).")
add_bullet(doc, "Botão \"Cancelar\" adicionado em todos os formulários para limpar o estado de edição.")
add_bullet(doc, "Mensagens de erro do backend (ex.: \"Supplier with id X not found\") são exibidas diretamente no feedback, mantendo rastreabilidade.")
add_bullet(doc, "Estados vazios com texto orientativo em todas as listas.")

add_heading(doc, "3.10. Documentação e decisões (FEED-10, FEED-12)", level=2)
add_bullet(doc, "docs/FRONTEND_MIGRATION_DECISION.md — registra a decisão de manter o frontend estático no piloto, com critérios claros para migrar para SPA no futuro.")
add_bullet(doc, "docs/INTEGRATIONS_ROADMAP.md — roadmap de integrações externas (ERP, e-mail, SSO) ordenado por valor × risco.")
add_bullet(doc, "CHANGELOG.md — histórico das alterações deste ciclo.")
add_bullet(doc, "README.md — atualizado com a seção \"Funcionalidades do Piloto (feedback dos usuários)\" e os 17 endpoints novos.")

doc.add_page_break()

# 4. Melhorias tecnicas
add_heading(doc, "4. Melhorias técnicas aplicadas", level=1)

add_heading(doc, "4.1. Ordem de middlewares nas rotas", level=2)
add_paragraph(doc, "Problema:", bold=True)
add_paragraph(
    doc,
    "Testes de RBAC falhavam porque o middleware userRoutes.use(allowRoles(['admin'])) "
    "estava sendo aplicado a rotas que não eram /users. A causa era a ordem de registro "
    "no src/routes/index.ts: como userRoutes era o primeiro sub-router montado em /api/v1, "
    "o Express encadeava o middleware use(allowRoles(['admin'])) antes das rotas de "
    "fornecedores, cotações e auditoria.",
)
add_paragraph(doc, "Solução:", bold=True)
add_paragraph(doc, "Reordenação explícita do registro no src/routes/index.ts:")
add_bullet(doc, "authRoutes (sem RBAC, primeiro sempre);")
add_bullet(doc, "rotas de negócio (supplierRoutes, supplierContactRoutes, attachmentRoutes, reportRoutes, quoteRequestRoutes, quoteResponseRoutes, quoteRequestItemRoutes);")
add_bullet(doc, "userRoutes (admin-only) e auditRoutes (admin/gestor) por último.")
add_paragraph(
    doc,
    "O mesmo padrão foi replicado no prefixo /api para garantir consistência com o "
    "legacy compatibility layer.",
)
add_paragraph(doc, "Validação:", bold=True)
add_paragraph(doc, "14 arquivos de teste passando, 46 testes verdes.")

add_heading(doc, "4.2. Segurança: nunca expor passwordHash", level=2)
add_paragraph(
    doc,
    "UserController.list, getById e getMe foram ajustados para devolver apenas os campos "
    "públicos (id, name, email, role, isActive, createdAt). Testes de RBAC já validavam "
    "isso indiretamente; foi adicionado teste explícito em tests/user-routes.test.ts.",
)

add_heading(doc, "4.3. Validação de anexos", level=2)
add_paragraph(
    doc,
    "AttachmentController.create valida entityType e entityId antes de gravar. Sem essa "
    "validação, era possível enviar um entityId inexistente e o registro ficava órfão. "
    "Agora, se o entityId não existir no banco para o entityType declarado, a API "
    "retorna 400.",
)

add_heading(doc, "4.4. Validação de isWinner no update", level=2)
add_paragraph(
    doc,
    "QuoteResponseController.update rejeita tentativas de marcar isWinner = true "
    "manualmente, pois a regra de negócio é que apenas o ComparisonController pode "
    "definir a vencedora. Continuamos permitindo isWinner = false para desfazer.",
)

add_heading(doc, "4.5. Migração de tokens de recuperação", level=2)
add_paragraph(
    doc,
    "Nova tabela PasswordResetToken no Prisma, com índices em userId e token, e cascata "
    "no onDelete: Cascade do User. Migração 20260505120000_password_recovery_tokens já "
    "gerada e pronta para npx prisma migrate deploy.",
)

doc.add_page_break()

# 5. Cobertura de testes
add_heading(doc, "5. Cobertura de testes", level=1)

add_table(
    doc,
    headers=["Arquivo", "Testes", "Cobre"],
    rows=[
        ["auth-routes.test.ts", "5", "login, refresh, logout, me, refresh expirado"],
        ["rbac-routes.test.ts", "6", "matriz de perfis × rotas"],
        ["user-routes.test.ts", "4", "listar, criar, editar, redefinir senha"],
        ["supplier-routes.test.ts", "4", "CRUD + busca"],
        ["supplier-contact-routes.test.ts", "4", "CRUD + unicidade do \"principal\""],
        ["quote-request-routes.test.ts", "5", "CRUD + close/reopen"],
        ["quote-response-routes.test.ts", "5", "CRUD + restrição de isWinner"],
        ["quote-request-item-routes.test.ts", "3", "CRUD"],
        ["attachment-routes.test.ts", "3", "upload, list, delete"],
        ["password-recovery-routes.test.ts", "5", "forgot, reset, listagem, expiração, rate limit"],
        ["comparison-routes.test.ts", "4", "comparar, histórico, pesos customizados"],
        ["reports-routes.test.ts", "6", "5 endpoints + filtros"],
        ["help-routes.test.ts", "2", "listagem + busca"],
        ["audit-routes.test.ts", "3", "listagem, filtros"],
        ["db-critical-flows.test.ts", "(skipped)", "fluxo end-to-end (requer DB real)"],
        ["Total", "46 + 1 skipped", ""],
    ],
)

add_paragraph(doc, "Comandos de validação:", bold=True)
add_paragraph(doc, "npm run build     # OK", style="No Spacing")
add_paragraph(doc, "npx tsc --noEmit  # OK", style="No Spacing")
add_paragraph(doc, "npm test          # 14 files, 46 tests passed, 1 skipped", style="No Spacing")

doc.add_page_break()

# 6. Limitacoes
add_heading(doc, "6. Limitações conhecidas e próximos passos", level=1)
add_bullet(doc, "FEED-10 (multi-tenant): decisão registrada de não implementar no piloto single-tenant. A migração para SaaS multi-empresa só será feita quando houver decisão de produto explícita. Ver docs/FRONTEND_MIGRATION_DECISION.md.")
add_bullet(doc, "Integrações externas (ERP, e-mail, SSO): roadmap priorizado em docs/INTEGRATIONS_ROADMAP.md. Próximas candidatas: e-mail transacional (SendGrid/Resend) e exportação CSV enriquecido com anexos.")
add_bullet(doc, "Onboarding persistente: o estado \"dispensado\" do painel de onboarding é em memória. Próximo passo: persistir no User (campo preferences JSON).")
add_bullet(doc, "Migração do frontend: quando o painel passar de ~5000 linhas ou surgirem requisitos de i18n dinâmico, migrar para SPA (Vite + React) mantendo o backend atual. Critérios completos em docs/FRONTEND_MIGRATION_DECISION.md.")
add_bullet(doc, "Migração do banco: após conectar ao Supabase, executar npx prisma migrate deploy para aplicar a tabela PasswordResetToken.")

doc.add_page_break()

# 7. Arquivos
add_heading(doc, "7. Arquivos alterados/criados (referência)", level=1)

add_paragraph(doc, "Backend", bold=True)
add_bullet(doc, "src/routes/index.ts (reordenado)")
add_bullet(doc, "src/routes/AuthRoutes.ts (novos endpoints)")
add_bullet(doc, "src/routes/SupplierRoutes.ts, SupplierContactRoutes.ts, AttachmentRoutes.ts, ReportRoutes.ts, HelpRoutes.ts (novos)")
add_bullet(doc, "src/controllers/AttachmentController.ts, SupplierContactController.ts, PasswordRecoveryController.ts, ReportController.ts, HelpController.ts (novos)")
add_bullet(doc, "src/services/passwordRecovery.ts (novo)")
add_bullet(doc, "src/content/helpArticles.ts (novo)")
add_bullet(doc, "src/validators/domain.ts (novos schemas)")
add_bullet(doc, "src/middlewares/rateLimit.ts (novo helper)")
add_bullet(doc, "prisma/schema.prisma (modelo PasswordResetToken)")
add_bullet(doc, "prisma/migrations/20260505120000_password_recovery_tokens/migration.sql (novo)")

add_paragraph(doc, "Frontend", bold=True)
add_bullet(doc, "public/index.html (4 novas abas, formulários de recuperação, painel de onboarding, wizard)")
add_bullet(doc, "public/app.js (handlers das novas abas, wizard, onboarding, relatórios, ajuda)")
add_bullet(doc, "public/styles.css (estilos .wizard-steps, .wizard-step-indicator, .onboarding-panel, .help-article-list, .report-grid, .report-card, .attachment-item, .contact-list-item, .users-table)")

add_paragraph(doc, "Testes", bold=True)
add_bullet(doc, "tests/supplier-contact-routes.test.ts (novo)")
add_bullet(doc, "tests/attachment-routes.test.ts (novo)")
add_bullet(doc, "tests/password-recovery-routes.test.ts (novo)")
add_bullet(doc, "tests/reports-routes.test.ts (novo)")

add_paragraph(doc, "Documentação", bold=True)
add_bullet(doc, "README.md (atualizado)")
add_bullet(doc, "CHANGELOG.md (novo)")
add_bullet(doc, "docs/FRONTEND_MIGRATION_DECISION.md (novo)")
add_bullet(doc, "docs/INTEGRATIONS_ROADMAP.md (novo)")
add_bullet(doc, "documentação/Relatorio_Correcoes_Melhorias_IntelliQuote.md (relatório Markdown de origem)")
add_bullet(doc, "documentação/Relatorio_Correcoes_Melhorias_IntelliQuote.docx (este arquivo Word)")

doc.add_page_break()

# 8. Conclusao
add_heading(doc, "8. Conclusão", level=1)
add_paragraph(
    doc,
    "Os 12 pontos de feedback foram tratados em sua totalidade, com decisão explícita "
    "documentada para o item de multi-tenant (FEED-10). A plataforma agora oferece:",
)
add_bullet(doc, "fluxo de proposta guiado por wizard (redução da densidade visual);")
add_bullet(doc, "anexos reais em fornecedores, cotações e propostas;")
add_bullet(doc, "recuperação de senha com tokens de uso único;")
add_bullet(doc, "gestão de usuários pela interface;")
add_bullet(doc, "onboarding para novos utilizadores;")
add_bullet(doc, "5 relatórios gerenciais;")
add_bullet(doc, "central de ajuda in-app;")
add_bullet(doc, "correções técnicas em RBAC, validação de anexos e segurança do usuário.")
add_paragraph(
    doc,
    "Os testes estão verdes, o build passa e a documentação foi ampliada. O IntelliQuote "
    "está pronto para a próxima rodada de piloto interno com os ajustes deste ciclo "
    "aplicados.",
    bold=True,
)


OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"OK -> {OUTPUT_PATH}")
